from flask import Flask, render_template, send_from_directory, request
from docx import Document
from read import readDocument
app = Flask(__name__)


@app.route('/')
def hello():
    return render_template('index.html')

@app.route('/results', methods=['POST'])
def results():
    protocolFile = request.files['protocol']
    document = Document(protocolFile)
    rows, results = readDocument(document);
    #print(document.paragraphs[0].text)
    return render_template('results.html', rows=rows, results=results)

if __name__ == '__main__':
    app.run()
