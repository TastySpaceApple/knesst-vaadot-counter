from docx import Document
import re


document = Document('protocol53knesset22.docx')


def readDocument(document):
    isSpeech = False
    isChair = False
    isMK = False
    speakerName = ""

    chairWords = 0
    mkWords = 0
    guestWords = 0
    rows = []

    for p in document.paragraphs:
        if p.text.strip() == "": continue
        #x = re.search("<<([^>]+)>>(.*)<<[^>]+>>\s*", p.text); # title regex
        style = p.style.name
        x = re.search("<<[^>]+>>(.*)<<[^>]+>>\s*", p.text);
        print(style)
        if style == u"יור":
            speakerName = p.text if x == None else x[1]
            isSpeech = True
            isChair = True
            isMK = False
        elif style == u"דובר" or style == u"דובר_המשך" or style == u"קריאה" or style == u"אורח":
            speakerName = p.text if x == None else x[1]
            isSpeech = True
            isChair = False
            isMK = re.match("[^\(]+\([^\)]+\)\:\s*$", speakerName) != None
        elif isSpeech:
            row = {"isMK": isMK, "isChair": isChair, "chairWords": 0, "mkWords": 0, "guestWords":0, "speakerName": speakerName, "text":p.text.strip()}
            text_length = len(row["text"].split(" "))
            if isChair:
                row["chairWords"] = text_length
                chairWords += text_length
                print(p.text)
            elif isMK:
                row["mkWords"] = text_length
                mkWords += text_length
            else:
                row["guestWords"] = text_length
                guestWords += text_length
            rows.append(row)

    results = {"chairWords": chairWords, "mkWords":mkWords, "guestWords": guestWords, "total": chairWords + mkWords + guestWords}
    results["chairWordsPercentage"] = chairWords * 100 / results["total"]
    results["mkWordsPercentage"] = mkWords * 100 / results["total"]
    results["guestWordsPercentage"] = guestWords * 100 / results["total"]
    return (rows, results)
    #print(chairWords * 100 / (otherWords + chairWords))
